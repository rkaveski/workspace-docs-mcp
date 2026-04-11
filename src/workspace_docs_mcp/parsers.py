from __future__ import annotations

import csv
import os
import shutil
from dataclasses import dataclass
from pathlib import Path

from .models import IMAGE_EXTENSIONS, ParsedSegment


class ParserError(RuntimeError):
    pass


@dataclass(frozen=True)
class OCRRuntime:
    enabled: bool
    available: bool
    reason: str | None
    lang: str
    timeout_seconds: int


def parse_document(path: Path) -> list[ParsedSegment]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        return _parse_markdown(path)
    if suffix == ".txt":
        return [ParsedSegment(text=path.read_text(encoding="utf-8", errors="ignore"))]
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".csv":
        return _parse_csv(path)
    if suffix == ".xlsx":
        return _parse_xlsx(path)
    if suffix in IMAGE_EXTENSIONS:
        return _parse_image_ocr(path)
    raise ParserError(f"Unsupported file type: {path}")


def get_parser_capabilities() -> dict:
    ocr = _get_ocr_runtime()
    return {
        "image_ocr": {
            "enabled": ocr.enabled,
            "available": ocr.available,
            "reason": ocr.reason,
            "lang": ocr.lang,
            "timeout_seconds": ocr.timeout_seconds,
        },
        "csv": True,
        "xlsx": True,
        "pdf_text": True,
    }


def _parse_markdown(path: Path) -> list[ParsedSegment]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    segments: list[ParsedSegment] = []
    current_title: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        content = "\n".join(buffer).strip()
        buffer.clear()
        if content:
            segments.append(ParsedSegment(text=content, section_title=current_title))

    for line in lines:
        if line.lstrip().startswith("#"):
            flush()
            current_title = line.lstrip("#").strip() or current_title
            continue
        buffer.append(line)

    flush()

    if not segments:
        segments.append(ParsedSegment(text=text))
    return segments


def _parse_pdf(path: Path) -> list[ParsedSegment]:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - dependency optional at runtime
        raise ParserError("pypdf is not installed") from exc

    reader = PdfReader(str(path))
    segments: list[ParsedSegment] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            segments.append(ParsedSegment(text=text, page_number=i))

    if not segments:
        segments.append(ParsedSegment(text=""))
    return segments


def _parse_docx(path: Path) -> list[ParsedSegment]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency optional at runtime
        raise ParserError("python-docx is not installed") from exc

    doc = Document(str(path))
    segments: list[ParsedSegment] = []
    current_title: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        content = "\n".join(buffer).strip()
        buffer.clear()
        if content:
            segments.append(ParsedSegment(text=content, section_title=current_title))

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style_name = (p.style.name or "").lower() if p.style else ""
        if style_name.startswith("heading"):
            flush()
            current_title = txt
            continue
        buffer.append(txt)

    flush()

    if not segments:
        whole = "\n".join((p.text or "") for p in doc.paragraphs).strip()
        segments.append(ParsedSegment(text=whole))

    return segments


def _parse_csv(path: Path) -> list[ParsedSegment]:
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        return _parse_tabular_rows(reader, section_title="sheet:csv")


def _parse_xlsx(path: Path) -> list[ParsedSegment]:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover - dependency optional at runtime
        raise ParserError("openpyxl is not installed") from exc

    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    out: list[ParsedSegment] = []
    for sheet in wb.worksheets:
        rows = sheet.iter_rows(values_only=True)
        out.extend(_parse_tabular_rows(rows, section_title=f"sheet:{sheet.title}"))
    wb.close()
    return out


def _parse_tabular_rows(rows, *, section_title: str) -> list[ParsedSegment]:
    max_rows = _env_int("WORKSPACE_DOCS_MAX_ROWS_PER_TABLE_FILE", 25_000, minimum=1, maximum=1_000_000)
    max_cell_chars = _env_int("WORKSPACE_DOCS_MAX_CELL_CHARS", 500, minimum=1, maximum=10_000)

    segments: list[ParsedSegment] = []
    header: list[str] | None = None
    data_rows = 0

    for raw in rows:
        cells = [_normalize_cell(v, max_cell_chars) for v in (raw or [])]
        if not any(cells):
            continue

        if header is None:
            header = [c if c else f"column_{idx + 1}" for idx, c in enumerate(cells)]
            continue

        data_rows += 1
        if data_rows > max_rows:
            segments.append(
                ParsedSegment(
                    text=f"TRUNCATED: indexed first {max_rows} rows only",
                    section_title=section_title,
                )
            )
            break

        values = list(cells)
        if len(values) < len(header):
            values.extend([""] * (len(header) - len(values)))
        elif len(values) > len(header):
            for i in range(len(header), len(values)):
                header.append(f"column_{i + 1}")

        pairs = [f"{header[i]}={values[i]}" for i in range(len(values)) if values[i] != ""]
        if not pairs:
            continue

        segments.append(ParsedSegment(text=" | ".join(pairs), section_title=section_title))

    return segments


def _parse_image_ocr(path: Path) -> list[ParsedSegment]:
    runtime = _get_ocr_runtime()
    if not runtime.enabled:
        return []
    if not runtime.available:
        raise ParserError(f"Image OCR unavailable: {runtime.reason}")

    try:
        from PIL import Image
        import pytesseract
    except Exception as exc:
        raise ParserError("Image OCR dependencies are not installed") from exc

    try:
        with Image.open(path) as img:
            text = (
                pytesseract.image_to_string(
                    img,
                    lang=runtime.lang,
                    timeout=runtime.timeout_seconds,
                )
                or ""
            ).strip()
    except RuntimeError as exc:
        raise ParserError(f"Image OCR timeout after {runtime.timeout_seconds}s: {exc}") from exc
    except Exception as exc:
        raise ParserError(f"Image OCR failed: {exc}") from exc

    if not text:
        return []
    return [ParsedSegment(text=text, section_title="image_ocr")]


def _get_ocr_runtime() -> OCRRuntime:
    enabled = _env_bool("WORKSPACE_DOCS_ENABLE_IMAGE_OCR", False)
    lang = os.getenv("WORKSPACE_DOCS_OCR_LANG", "eng").strip() or "eng"
    timeout_seconds = _env_int("WORKSPACE_DOCS_OCR_TIMEOUT_SECONDS", 15, minimum=1, maximum=600)

    if not enabled:
        return OCRRuntime(
            enabled=False,
            available=False,
            reason="disabled via WORKSPACE_DOCS_ENABLE_IMAGE_OCR",
            lang=lang,
            timeout_seconds=timeout_seconds,
        )

    try:
        import pytesseract
    except Exception:
        return OCRRuntime(
            enabled=True,
            available=False,
            reason="pytesseract is not installed",
            lang=lang,
            timeout_seconds=timeout_seconds,
        )

    try:
        from PIL import Image  # noqa: F401
    except Exception:
        return OCRRuntime(
            enabled=True,
            available=False,
            reason="Pillow is not installed",
            lang=lang,
            timeout_seconds=timeout_seconds,
        )

    binary = shutil.which("tesseract")
    if not binary:
        return OCRRuntime(
            enabled=True,
            available=False,
            reason="tesseract binary not found in PATH",
            lang=lang,
            timeout_seconds=timeout_seconds,
        )

    try:
        _ = pytesseract.get_tesseract_version()
    except Exception as exc:
        return OCRRuntime(
            enabled=True,
            available=False,
            reason=f"tesseract unavailable: {exc}",
            lang=lang,
            timeout_seconds=timeout_seconds,
        )

    return OCRRuntime(enabled=True, available=True, reason=None, lang=lang, timeout_seconds=timeout_seconds)


def _env_bool(name: str, default: bool) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    val = raw.strip().lower()
    if val in {"1", "true", "yes", "on"}:
        return True
    if val in {"0", "false", "no", "off"}:
        return False
    return default


def _env_int(name: str, default: int, *, minimum: int, maximum: int) -> int:
    raw = os.getenv(name)
    if raw is None:
        return default
    try:
        value = int(raw)
    except Exception:
        return default
    return min(maximum, max(minimum, value))


def _normalize_cell(value, max_chars: int) -> str:
    if value is None:
        return ""
    text = str(value).strip().replace("\n", " ").replace("\r", " ")
    if len(text) <= max_chars:
        return text
    if max_chars <= 3:
        return text[:max_chars]
    return text[: max_chars - 3].rstrip() + "..."
